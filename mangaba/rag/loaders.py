"""
Document loaders for Mangaba AI RAG pipeline.

Every loader exposes the same contract — ``load() -> List[Document]`` — and
every produced :class:`~mangaba.rag.document.Document` carries a ``source``
metadata key identifying where the content came from (file path or URL).
Loaders that split a file into several documents add their own positional
metadata (``page``, ``row``, ``sheet``, ``index``).

Optional third-party dependencies (``pypdf``, ``python-docx``, ``openpyxl``)
are imported lazily inside the loader that needs them, so importing this
module never requires them.
"""

from __future__ import annotations

import csv
import json
import logging
from pathlib import Path
from typing import Any, Dict, List, Optional, Type, Union

from mangaba.rag.document import Document

log = logging.getLogger(__name__)


class TextLoader:
    """Load a plain-text file as a single Document.

    Example::

        docs = TextLoader("notes.txt").load()
        print(docs[0].metadata["source"])
    """

    def __init__(self, file_path: str, encoding: str = "utf-8") -> None:
        self.file_path = file_path
        self.encoding = encoding

    def load(self) -> List[Document]:
        path = Path(self.file_path)
        text = path.read_text(encoding=self.encoding)
        return [Document(content=text, metadata={"source": str(path)})]


class CSVLoader:
    """Load a CSV file — each row becomes a Document.

    Example::

        docs = CSVLoader("faq.csv", content_columns=["question", "answer"]).load()
    """

    def __init__(
        self,
        file_path: str,
        content_columns: Optional[List[str]] = None,
        encoding: str = "utf-8",
    ) -> None:
        self.file_path = file_path
        self.content_columns = content_columns
        self.encoding = encoding

    def load(self) -> List[Document]:
        path = Path(self.file_path)
        docs: List[Document] = []
        with path.open(encoding=self.encoding, newline="") as fh:
            reader = csv.DictReader(fh)
            for i, row in enumerate(reader):
                if self.content_columns:
                    text = " | ".join(str(row.get(c, "")) for c in self.content_columns)
                else:
                    text = " | ".join(str(v) for v in row.values())
                meta: Dict[str, str] = {"source": str(path), "row": str(i)}
                docs.append(Document(content=text, metadata=meta))
        return docs


class WebPageLoader:
    """Load a web page as a Document (requires requests + beautifulsoup4).

    Example::

        docs = WebPageLoader("https://example.com").load()
    """

    def __init__(self, url: str) -> None:
        self.url = url

    def load(self) -> List[Document]:
        import requests  # already in dependencies
        resp = requests.get(self.url, timeout=15)
        resp.raise_for_status()
        try:
            from bs4 import BeautifulSoup  # type: ignore
            soup = BeautifulSoup(resp.text, "html.parser")
            # Remove script/style
            for tag in soup(["script", "style"]):
                tag.decompose()
            text = soup.get_text(separator="\n", strip=True)
        except ImportError:
            text = resp.text
        return [Document(content=text, metadata={"source": self.url})]


class PDFLoader:
    """Load a PDF file using ``pypdf``.

    By default one Document is produced per page (metadata ``page`` is
    1-based). Set ``split_pages=False`` to get a single Document with all
    pages joined by ``page_separator``.

    Example::

        docs = PDFLoader("report.pdf").load()
        print(docs[0].metadata)   # {'source': 'report.pdf', 'page': 1, 'total_pages': 12}

        whole = PDFLoader("report.pdf", split_pages=False).load()[0]
    """

    def __init__(
        self,
        file_path: str,
        split_pages: bool = True,
        page_separator: str = "\n\n",
        password: Optional[str] = None,
    ) -> None:
        self.file_path = file_path
        self.split_pages = split_pages
        self.page_separator = page_separator
        self.password = password

    def load(self) -> List[Document]:
        try:
            from pypdf import PdfReader  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on env
            raise ImportError(
                "pypdf package is required to load PDF files. "
                "Install with: pip install mangaba[documents]"
            ) from exc

        path = Path(self.file_path)
        reader = PdfReader(str(path))
        if getattr(reader, "is_encrypted", False) and self.password is not None:
            reader.decrypt(self.password)

        pages: List[str] = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:  # noqa: BLE001 - a broken page must not kill the load
                log.warning("PDFLoader: failed to extract text from a page of %s: %s", path, exc)
                pages.append("")

        total = len(pages)
        if not self.split_pages:
            text = self.page_separator.join(p for p in pages if p.strip())
            return [Document(content=text, metadata={"source": str(path), "total_pages": total})]

        docs: List[Document] = []
        for i, text in enumerate(pages):
            if not text.strip():
                continue
            meta: Dict[str, Any] = {"source": str(path), "page": i + 1, "total_pages": total}
            docs.append(Document(content=text, metadata=meta))
        return docs


class DOCXLoader:
    """Load a Word ``.docx`` file using ``python-docx``.

    Paragraphs are joined with newlines; table cells are appended row by row
    (``" | "`` separated) when ``include_tables`` is True.

    Example::

        docs = DOCXLoader("contract.docx").load()
        print(docs[0].metadata["paragraphs"])
    """

    def __init__(self, file_path: str, include_tables: bool = True) -> None:
        self.file_path = file_path
        self.include_tables = include_tables

    def load(self) -> List[Document]:
        try:
            import docx  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on env
            raise ImportError(
                "python-docx package is required to load DOCX files. "
                "Install with: pip install mangaba[documents]"
            ) from exc

        path = Path(self.file_path)
        document = docx.Document(str(path))

        parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]
        paragraph_count = len(parts)

        table_count = 0
        if self.include_tables:
            for table in document.tables:
                table_count += 1
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    if any(cells):
                        parts.append(" | ".join(cells))

        meta: Dict[str, Any] = {
            "source": str(path),
            "paragraphs": paragraph_count,
            "tables": table_count,
        }
        return [Document(content="\n".join(parts), metadata=meta)]


class ExcelLoader:
    """Load an Excel workbook (``.xlsx``/``.xlsm``) using ``openpyxl``.

    Handles multiple sheets. With the default ``mode="sheet"`` every sheet
    becomes one Document; with ``mode="row"`` every non-empty row becomes its
    own Document (mirroring :class:`CSVLoader`).

    Example::

        # one Document per sheet
        docs = ExcelLoader("sales.xlsx").load()
        print([d.metadata["sheet"] for d in docs])

        # one Document per row, only from a given sheet
        rows = ExcelLoader("sales.xlsx", mode="row", sheet_names=["2026"]).load()
    """

    def __init__(
        self,
        file_path: str,
        sheet_names: Optional[List[str]] = None,
        mode: str = "sheet",
        separator: str = " | ",
    ) -> None:
        if mode not in ("sheet", "row"):
            raise ValueError(f"mode must be 'sheet' or 'row', got {mode!r}")
        self.file_path = file_path
        self.sheet_names = sheet_names
        self.mode = mode
        self.separator = separator

    def load(self) -> List[Document]:
        try:
            from openpyxl import load_workbook  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on env
            raise ImportError(
                "openpyxl package is required to load Excel files. "
                "Install with: pip install mangaba[documents]"
            ) from exc

        path = Path(self.file_path)
        workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
        try:
            wanted = self.sheet_names or list(workbook.sheetnames)
            docs: List[Document] = []
            for name in wanted:
                if name not in workbook.sheetnames:
                    log.warning("ExcelLoader: sheet %r not found in %s — skipping", name, path)
                    continue
                sheet = workbook[name]
                rows = [self._format_row(r) for r in sheet.iter_rows(values_only=True)]
                rows = [r for r in rows if r.strip()]

                if self.mode == "sheet":
                    if not rows:
                        continue
                    meta: Dict[str, Any] = {
                        "source": str(path),
                        "sheet": name,
                        "rows": len(rows),
                    }
                    docs.append(Document(content="\n".join(rows), metadata=meta))
                else:
                    for i, text in enumerate(rows):
                        row_meta: Dict[str, Any] = {
                            "source": str(path),
                            "sheet": name,
                            "row": i,
                        }
                        docs.append(Document(content=text, metadata=row_meta))
            return docs
        finally:
            workbook.close()

    def _format_row(self, row: Any) -> str:
        return self.separator.join("" if cell is None else str(cell) for cell in row)


class JSONLoader:
    """Load a JSON file, optionally selecting fields with a dot-path.

    ``jq_like`` is a deliberately tiny, dependency-free subset of the ``jq``
    syntax: dot-separated keys, where a ``[]`` suffix means "iterate this
    list". For example ``"data.items[].text"`` walks into ``data``, then
    ``items``, expands the list, and picks ``text`` from every element.
    Missing keys are skipped silently rather than raising.

    Without ``jq_like``: a top-level list produces one Document per element,
    anything else produces a single Document with the pretty-printed JSON.
    Selected values that are not strings are serialized with ``json.dumps``.

    Example::

        # {"data": {"items": [{"text": "a"}, {"text": "b"}]}}
        docs = JSONLoader("payload.json", jq_like="data.items[].text").load()
        [d.content for d in docs]        # ['a', 'b']
        docs[0].metadata                 # {'source': ..., 'index': 0, 'json_path': 'data.items[].text'}
    """

    def __init__(
        self,
        file_path: str,
        jq_like: Optional[str] = None,
        encoding: str = "utf-8",
        ensure_ascii: bool = False,
    ) -> None:
        self.file_path = file_path
        self.jq_like = jq_like
        self.encoding = encoding
        self.ensure_ascii = ensure_ascii

    def load(self) -> List[Document]:
        path = Path(self.file_path)
        data = json.loads(path.read_text(encoding=self.encoding))

        if self.jq_like:
            values = self._select(data, self.jq_like)
        elif isinstance(data, list):
            values = list(data)
        else:
            values = [data]

        single = self.jq_like is None and not isinstance(data, list)
        docs: List[Document] = []
        for i, value in enumerate(values):
            text = self._stringify(value, pretty=single)
            if not text.strip():
                continue
            meta: Dict[str, Any] = {"source": str(path)}
            if not single:
                meta["index"] = i
            if self.jq_like:
                meta["json_path"] = self.jq_like
            docs.append(Document(content=text, metadata=meta))
        return docs

    def _stringify(self, value: Any, pretty: bool = False) -> str:
        if isinstance(value, str):
            return value
        indent = 2 if pretty else None
        return json.dumps(value, ensure_ascii=self.ensure_ascii, indent=indent)

    @staticmethod
    def _select(data: Any, path: str) -> List[Any]:
        """Resolve a dot-path such as ``a.b[].c`` into a flat list of values."""
        values: List[Any] = [data]
        for raw_token in path.split("."):
            token = raw_token.strip()
            if not token:
                continue
            expansions = 0
            while token.endswith("[]"):
                token = token[:-2]
                expansions += 1

            picked: List[Any] = []
            for value in values:
                if not token:
                    picked.append(value)
                elif isinstance(value, dict) and token in value:
                    picked.append(value[token])
                else:
                    log.debug("JSONLoader: key %r not found while resolving %r", token, path)

            for _ in range(expansions):
                flattened: List[Any] = []
                for value in picked:
                    if isinstance(value, list):
                        flattened.extend(value)
                    else:
                        flattened.append(value)
                picked = flattened

            values = picked
        return values


#: Maps a lowercase file extension to the loader class handling it.
EXTENSION_LOADERS: Dict[str, Type[Any]] = {
    ".txt": TextLoader,
    ".md": TextLoader,
    ".markdown": TextLoader,
    ".rst": TextLoader,
    ".log": TextLoader,
    ".csv": CSVLoader,
    ".json": JSONLoader,
    ".pdf": PDFLoader,
    ".docx": DOCXLoader,
    ".xlsx": ExcelLoader,
    ".xlsm": ExcelLoader,
}


class DirectoryLoader:
    """Walk a directory and dispatch every file to the matching loader.

    Files whose extension has no registered loader are skipped. When
    ``silent_errors`` is True (the default) a file that fails to load — for
    instance because an optional dependency is missing — is logged and
    skipped instead of aborting the whole walk.

    Example::

        docs = DirectoryLoader("./docs", glob="*.md", recursive=True).load()
        print({d.metadata["source"] for d in docs})
    """

    def __init__(
        self,
        path: str,
        glob: str = "*",
        recursive: bool = True,
        silent_errors: bool = True,
        exclude_hidden: bool = True,
        extension_loaders: Optional[Dict[str, Type[Any]]] = None,
        loader_kwargs: Optional[Dict[str, Dict[str, Any]]] = None,
    ) -> None:
        self.path = path
        self.glob = glob
        self.recursive = recursive
        self.silent_errors = silent_errors
        self.exclude_hidden = exclude_hidden
        self.extension_loaders = dict(extension_loaders or EXTENSION_LOADERS)
        self.loader_kwargs = loader_kwargs or {}

    def load(self) -> List[Document]:
        root = Path(self.path)
        if not root.is_dir():
            raise NotADirectoryError(f"{root} is not a directory")

        docs: List[Document] = []
        for file_path in self._iter_files(root):
            suffix = file_path.suffix.lower()
            loader_cls = self.extension_loaders.get(suffix)
            if loader_cls is None:
                log.debug("DirectoryLoader: no loader for %s — skipping", file_path)
                continue
            kwargs = self.loader_kwargs.get(suffix, {})
            try:
                docs.extend(loader_cls(str(file_path), **kwargs).load())
            except Exception as exc:  # noqa: BLE001 - one bad file must not kill the walk
                if not self.silent_errors:
                    raise
                log.warning("DirectoryLoader: failed to load %s: %s", file_path, exc)
        return docs

    def _iter_files(self, root: Path) -> List[Path]:
        matches = root.rglob(self.glob) if self.recursive else root.glob(self.glob)
        files = [p for p in matches if p.is_file()]
        if self.exclude_hidden:
            files = [p for p in files if not self._is_hidden(p, root)]
        return sorted(files)

    @staticmethod
    def _is_hidden(file_path: Path, root: Path) -> bool:
        try:
            relative = file_path.relative_to(root)
        except ValueError:  # pragma: no cover - defensive
            relative = file_path
        return any(part.startswith(".") for part in relative.parts)


def load_file(file_path: Union[str, Path], **kwargs: Any) -> List[Document]:
    """Load a single file with the loader registered for its extension.

    Example::

        docs = load_file("report.pdf")

    Raises:
        ValueError: If no loader is registered for the file extension.
    """
    path = Path(file_path)
    loader_cls = EXTENSION_LOADERS.get(path.suffix.lower())
    if loader_cls is None:
        raise ValueError(
            f"No loader registered for extension {path.suffix!r}. "
            f"Supported: {', '.join(sorted(EXTENSION_LOADERS))}"
        )
    return loader_cls(str(path), **kwargs).load()
